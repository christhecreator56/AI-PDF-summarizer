import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from typing import Dict, List, Any
import json
from datetime import datetime

class StudentFormatter:
    """
    Formats AI summaries into student-friendly outputs with learning aids.
    """
    
    def __init__(self):
        """Initialize the student formatter."""
        self.learning_aids = {
            'key_points': '🔑',
            'important': '⚠️',
            'example': '💡',
            'definition': '📖',
            'summary': '📝',
            'question': '❓',
            'tip': '💡'
        }
    
    def format_summary_for_students(self, summary_data: Dict[str, Any]) -> Dict[str, str]:
        """
        Format the AI summary into multiple student-friendly formats.
        
        Args:
            summary_data: Dictionary containing summary and metadata
            
        Returns:
            Dictionary with different formatted outputs
        """
        summary = summary_data.get('summary', '')
        readability_scores = summary_data.get('readability_scores', {})
        recommendations = summary_data.get('recommendations', [])
        
        # Create different formats
        formats = {
            'simple_text': self._create_simple_text(summary, readability_scores),
            'structured_text': self._create_structured_text(summary, recommendations),
            'markdown': self._create_markdown_format(summary, summary_data),
            'html': self._create_html_format(summary, summary_data),
            'json': self._create_json_format(summary_data)
        }
        
        return formats
    
    def _create_simple_text(self, summary: str, readability_scores: Dict[str, float]) -> str:
        """
        Create a simple text format with basic formatting.
        
        Args:
            summary: The summary text
            readability_scores: Readability metrics
            
        Returns:
            Formatted simple text
        """
        # Clean markdown for plain text
        text = re.sub(r'(\*\*|###\s?.*?\n|```.*?\n|```)', '', summary)

        text_content = f"""📚 STUDENT SUMMARY
{'='*50}

{text}

{'='*50}
📊 READING LEVEL ANALYSIS:
• Flesch Reading Ease: {readability_scores.get('flesch_reading_ease', 0):.1f}/100
• Grade Level: {readability_scores.get('flesch_kincaid_grade', 0):.1f}
• Words: {len(summary.split())}
• Sentences: {len(summary.split('.'))}

💡 TIP: This summary is designed to be easy to understand for students!
"""
        return text_content
    
    def _create_structured_text(self, summary: str, recommendations: List[str]) -> str:
        """
        Create a structured text format with sections and learning aids.
        
        Args:
            summary: The summary text
            recommendations: List of recommendations
            
        Returns:
            Structured text format
        """
        # Clean markdown for plain text
        summary = re.sub(r'(\*\*|###\s?.*?\n|```.*?\n|```)', '', summary)

        # Split summary into sentences
        sentences = [s.strip() for s in summary.split('.') if s.strip()]
        
        # Identify key sentences (first few and those with important words)
        key_words = ['important', 'key', 'main', 'primary', 'essential', 'crucial', 'significant']
        key_sentences = []
        
        for i, sentence in enumerate(sentences):
            if i < 2 or any(word in sentence.lower() for word in key_words):
                key_sentences.append(i)
        
        # Create structured format
        structured_text = f"""📖 COMPREHENSIVE SUMMARY
{'='*60}

🔑 KEY POINTS:
"""
        
        for i, sentence in enumerate(sentences):
            if i in key_sentences:
                structured_text += f"• {sentence}.\n"
        
        structured_text += f"""
📝 FULL SUMMARY:
"""
        
        for sentence in sentences:
            structured_text += f"• {sentence}.\n"
        
        if recommendations:
            structured_text += f"""
💡 LEARNING TIPS:
"""
            for rec in recommendations:
                structured_text += f"• {rec}\n"
        
        return structured_text
    
    def _create_markdown_format(self, summary: str, summary_data: Dict[str, Any]) -> str:
        """
        Create a markdown format with rich formatting.
        
        Args:
            summary: The summary text
            summary_data: Complete summary data
            
        Returns:
            Markdown formatted text
        """
        readability_scores = summary_data.get('readability_scores', {})
        recommendations = summary_data.get('recommendations', [])
        
        markdown_text = f"""# 📚 Student-Friendly Summary

{summary}

---

## 📊 Reading Analysis

| Metric | Score |
|--------|-------|
| **Flesch Reading Ease** | {readability_scores.get('flesch_reading_ease', 0):.1f}/100 |
| **Grade Level** | {readability_scores.get('flesch_kincaid_grade', 0):.1f} |
| **Gunning Fog Index** | {readability_scores.get('gunning_fog', 0):.1f} |
| **SMOG Index** | {readability_scores.get('smog_index', 0):.1f} |

**Word Count:** {summary_data.get('word_count', 0)}  
**Sentence Count:** {summary_data.get('sentence_count', 0)}  
**Compression Ratio:** {summary_data.get('compression_ratio', 0):.1%}

---

## 💡 Learning Tips

"""
        
        for rec in recommendations:
            markdown_text += f"- {rec}\n"
        
        markdown_text += f"""

---

*Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}*
"""
        
        return markdown_text
    
    def _create_html_format(self, summary: str, summary_data: Dict[str, Any]) -> str:
        """
        Create an HTML format with styling.
        
        Args:
            summary: The summary text
            summary_data: Complete summary data
            
        Returns:
            HTML formatted text
        """
        md = markdown.Markdown(extensions=['fenced_code', 'tables'])
        summary_html = md.convert(summary)
        
        readability_scores = summary_data.get('readability_scores', {})
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Student Summary</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2, h3 {{
            color: #34495e;
            margin-top: 30px;
        }}
        .summary {{
            background: #ecf0f1;
            padding: 20px;
            border-left: 4px solid #3498db;
            margin: 20px 0;
        }}
        pre, code {{
            background-color: #eee;
            border: 1px solid #999;
            display: block;
            padding: 10px;
            border-radius: 5px;
            font-family: 'Courier New', Courier, monospace;
            white-space: pre-wrap;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 12px;
            border: 1px solid #ddd;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
<div class="container">
    <h1>Student Summary</h1>
    <div class="summary">
        {summary_html}
    </div>
</div>
</body>
</html>
"""
        return html
    
    def _create_json_format(self, summary_data: Dict[str, Any]) -> str:
        """
        Create a JSON format of the summary data.
        
        Args:
            summary_data: Complete summary data
            
        Returns:
            JSON formatted string
        """
        # Ensure all data is serializable
        serializable_data = {}
        for key, value in summary_data.items():
            try:
                json.dumps(value)
                serializable_data[key] = value
            except (TypeError, OverflowError):
                serializable_data[key] = str(value)
        
        return json.dumps(serializable_data, indent=2)
    
    def _extract_key_concepts(self, text: str) -> str:
        """
        Extract key concepts from the text (simple regex for bolded items).
        
        Args:
            text: Input text with markdown
            
        Returns:
            A string of key concepts.
        """
        # Find all bolded items
        concepts = re.findall(r'\*\*(.*?)\*\*', text)
        if not concepts:
            return "No key concepts highlighted."
        
        # Return unique concepts
        unique_concepts = sorted(list(set(c.lower() for c in concepts)))
        return ", ".join(unique_concepts).title()

    def create_word_document(self, summary_data: Dict[str, Any]):
        """
        Create a Word document from the summary with formatting for all elements.

        Args:
            summary_data: Dictionary containing summary and metadata

        Returns:
            The created Word document object.
        """
        doc = Document()
        doc.add_heading('📚 Student-Friendly Summary', level=1)
        
        summary_text = summary_data.get('summary', '')
        
        # Process the main summary content which contains markdown
        self._process_doc_content(doc, summary_text)
            
        return doc

    def _process_doc_content(self, doc, content: str):
        """Process and add standard content (headings, paragraphs, code) to the doc."""
        # Split by major headings (###) to maintain structure
        parts = re.split(r'(### .*?\n)', content)
        
        # The first part is anything before the first heading
        if parts[0].strip():
            self._add_content_chunk_to_doc(doc, parts[0])

        # Process the rest of the parts
        for i in range(1, len(parts), 2):
            heading = parts[i].replace('###', '').strip()
            content_chunk = parts[i+1]
            
            doc.add_heading(heading, level=3)
            self._add_content_chunk_to_doc(doc, content_chunk)

    def _add_content_chunk_to_doc(self, doc, content: str):
        """Adds a chunk of content, parsing for code blocks and bold text."""
        # Split content by code blocks to handle them separately
        sub_parts = re.split(r'(```python\n.*?\n```)', content, flags=re.DOTALL)
        for part in sub_parts:
            if not part.strip():
                continue
            
            if part.startswith('```python'):
                code_text = part.replace('```python\n', '').replace('\n```', '').strip()
                p = doc.add_paragraph()
                p.add_run(code_text).font.name = 'Courier New'
            else:
                self._add_paragraph_with_bolding(doc, part)

    def _add_paragraph_with_bolding(self, doc, text: str):
        """Adds a paragraph to the document, handling **bold** markdown."""
        p = doc.add_paragraph()
        # Split by bold markers, keeping the markers
        sub_parts = re.split(r'(\*\*.*?\*\*)', text)
        for sub_part in sub_parts:
            if sub_part.startswith('**') and sub_part.endswith('**'):
                p.add_run(sub_part[2:-2]).bold = True
            elif sub_part:
                p.add_run(sub_part) 